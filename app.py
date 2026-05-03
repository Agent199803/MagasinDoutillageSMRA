import customtkinter as ctk
import sqlite3
import os
import sys
import shutil
from datetime import datetime
from tkinter import messagebox, filedialog
from docx import Document

if getattr(sys, "frozen", False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

DB_NAME = os.path.join(BASE_DIR, "magasin.db")
DEFAULT_PASSWORD = "1234"

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")


def connect_db():
    return sqlite3.connect(DB_NAME)


def ensure_column(cur, table, column, definition):
    cur.execute(f"PRAGMA table_info({table})")
    cols = [x[1] for x in cur.fetchall()]
    if column not in cols:
        cur.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def init_db():
    conn = connect_db()
    cur = conn.cursor()

    cur.execute("CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT)")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS technicians (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT UNIQUE NOT NULL,
        matricule TEXT,
        service TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS tools (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        serial TEXT UNIQUE NOT NULL,
        category TEXT,
        status TEXT DEFAULT 'Disponible'
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS movements (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        technician TEXT,
        tool_name TEXT,
        serial TEXT,
        action TEXT,
        date_time TEXT
    )
    """)

    ensure_column(cur, "technicians", "matricule", "TEXT")
    ensure_column(cur, "technicians", "service", "TEXT")
    ensure_column(cur, "tools", "category", "TEXT")
    ensure_column(cur, "tools", "status", "TEXT DEFAULT 'Disponible'")
    ensure_column(cur, "movements", "tool_name", "TEXT")
    ensure_column(cur, "movements", "serial", "TEXT")
    ensure_column(cur, "movements", "date_time", "TEXT")

    cur.execute("SELECT value FROM settings WHERE key='password'")
    if cur.fetchone() is None:
        cur.execute("INSERT INTO settings (key, value) VALUES ('password', ?)", (DEFAULT_PASSWORD,))

    conn.commit()
    conn.close()


def get_password():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT value FROM settings WHERE key='password'")
    row = cur.fetchone()
    conn.close()
    return row[0] if row else DEFAULT_PASSWORD


def set_password(value):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("UPDATE settings SET value=? WHERE key='password'", (value,))
    conn.commit()
    conn.close()


class Login(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Login")
        self.geometry("430x300")
        self.resizable(False, False)

        box = ctk.CTkFrame(self, corner_radius=18)
        box.pack(fill="both", expand=True, padx=25, pady=25)

        ctk.CTkLabel(
            box,
            text="MAGASIN D'OUTILLAGE SMRA",
            font=("Arial", 23, "bold")
        ).pack(pady=(35, 10))

        ctk.CTkLabel(
            box,
            text="Default password: 1234",
            font=("Arial", 13)
        ).pack(pady=5)

        self.entry = ctk.CTkEntry(
            box,
            placeholder_text="Password",
            show="*",
            height=42
        )
        self.entry.pack(fill="x", padx=35, pady=15)
        self.entry.bind("<Return>", lambda e: self.login())

        ctk.CTkButton(
            box,
            text="Enter",
            height=42,
            command=self.login
        ).pack(fill="x", padx=35, pady=10)

    def login(self):
        if self.entry.get() == get_password():
            self.destroy()
            app = App()
            app.mainloop()
        else:
            messagebox.showerror("Error", "Incorrect password")


class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Magasin D'outillage SMRA")
        self.geometry("1200x720")
        self.minsize(1000, 620)

        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.sidebar = ctk.CTkFrame(self, width=240, corner_radius=0, fg_color="#0B1F33")
        self.sidebar.grid(row=0, column=0, sticky="nsw")
        self.sidebar.grid_propagate(False)

        self.main = ctk.CTkFrame(self, corner_radius=0, fg_color="#EAF6FF")
        self.main.grid(row=0, column=1, sticky="nsew")
        self.main.grid_columnconfigure(0, weight=1)
        self.main.grid_rowconfigure(1, weight=1)

        self.sidebar_ui()
        self.home()

    def sidebar_ui(self):
        ctk.CTkLabel(
            self.sidebar,
            text="MAGASIN\nD'OUTILLAGE\nSMRA",
            font=("Arial", 23, "bold"),
            text_color="white",
            justify="center"
        ).pack(pady=(35, 20))

        buttons = [
            ("Accueil", self.home),
            ("Techniciens", self.technicians_page),
            ("Outils", self.tools_page),
            ("Journal", self.journal_page),
            ("Paramètres", self.settings_page),
        ]

        for text, command in buttons:
            ctk.CTkButton(
                self.sidebar,
                text=text,
                command=command,
                height=45,
                corner_radius=12,
                fg_color="#179BD7",
                hover_color="#0E6FA3",
                text_color="white",
                font=("Arial", 15, "bold"),
                anchor="w"
            ).pack(fill="x", padx=18, pady=7)

        ctk.CTkLabel(
            self.sidebar,
            text="Version 1.0\nMOHAMMED BELKEBIR\nABDELKARIM",
            text_color="#B8DDF2",
            font=("Arial", 12),
            justify="center"
        ).pack(side="bottom", pady=25)

    def clear_main(self):
        for widget in self.main.winfo_children():
            widget.destroy()

    def page_title(self, title, subtitle=""):
        header = ctk.CTkFrame(self.main, fg_color="#EAF6FF")
        header.grid(row=0, column=0, sticky="ew", padx=25, pady=(20, 10))

        ctk.CTkLabel(
            header,
            text=title,
            font=("Arial", 30, "bold"),
            text_color="#071B2C"
        ).pack(anchor="w")

        if subtitle:
            ctk.CTkLabel(
                header,
                text=subtitle,
                font=("Arial", 15),
                text_color="#36556B"
            ).pack(anchor="w", pady=(5, 0))

    def card(self, parent):
        return ctk.CTkFrame(parent, fg_color="white", corner_radius=18)

    def home(self):
        self.clear_main()
        self.page_title("Tableau de bord", "Gestion des sorties et retours d'outillage")

        body = ctk.CTkFrame(self.main, fg_color="#EAF6FF")
        body.grid(row=1, column=0, sticky="nsew", padx=25, pady=10)
        body.grid_columnconfigure(0, weight=1)
        body.grid_rowconfigure(2, weight=1)

        stats_frame = ctk.CTkFrame(body, fg_color="#EAF6FF")
        stats_frame.grid(row=0, column=0, sticky="ew")
        stats_frame.grid_columnconfigure((0, 1, 2), weight=1)

        conn = connect_db()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM technicians")
        tech_count = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM tools")
        tool_count = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM tools WHERE status='Sortie'")
        out_count = cur.fetchone()[0]
        conn.close()

        stats = [
            ("Techniciens", tech_count),
            ("Outils", tool_count),
            ("Outils sortis", out_count),
        ]

        for i, item in enumerate(stats):
            box = self.card(stats_frame)
            box.grid(row=0, column=i, sticky="nsew", padx=8, pady=8)

            ctk.CTkLabel(
                box,
                text=str(item[1]),
                font=("Arial", 34, "bold"),
                text_color="#179BD7"
            ).pack(pady=(18, 3))

            ctk.CTkLabel(
                box,
                text=item[0],
                font=("Arial", 15, "bold"),
                text_color="#071B2C"
            ).pack(pady=(0, 18))

        movement_box = self.card(body)
        movement_box.grid(row=1, column=0, sticky="ew", pady=10)
        movement_box.grid_columnconfigure((0, 1), weight=1)

        ctk.CTkLabel(
            movement_box,
            text="Nouvelle sortie d'outil",
            font=("Arial", 20, "bold"),
            text_color="#071B2C"
        ).grid(row=0, column=0, columnspan=3, sticky="w", padx=15, pady=(15, 5))

        tech_entry = ctk.CTkEntry(movement_box, placeholder_text="Rechercher technicien...", height=40)
        tech_entry.grid(row=1, column=0, padx=10, pady=10, sticky="ew")

        tool_entry = ctk.CTkEntry(movement_box, placeholder_text="Rechercher outil ou numéro de série...", height=40)
        tool_entry.grid(row=1, column=1, padx=10, pady=10, sticky="ew")

        selected_tech = {"name": None}
        selected_tool = {"name": None, "serial": None}

        suggestions_frame = ctk.CTkFrame(movement_box, fg_color="white")
        suggestions_frame.grid(row=2, column=0, columnspan=3, sticky="ew", padx=10, pady=(0, 10))
        suggestions_frame.grid_columnconfigure((0, 1), weight=1)

        tech_suggestions = ctk.CTkScrollableFrame(suggestions_frame, height=110, fg_color="#F4FAFF")
        tech_suggestions.grid(row=0, column=0, sticky="ew", padx=5, pady=5)

        tool_suggestions = ctk.CTkScrollableFrame(suggestions_frame, height=110, fg_color="#F4FAFF")
        tool_suggestions.grid(row=0, column=1, sticky="ew", padx=5, pady=5)

        def clear_suggestions(frame):
            for w in frame.winfo_children():
                w.destroy()

        def search_technicians(event=None):
            clear_suggestions(tech_suggestions)
            selected_tech["name"] = None

            keyword = tech_entry.get().strip()
            if not keyword:
                return

            conn = connect_db()
            cur = conn.cursor()
            like = f"%{keyword}%"
            cur.execute("""
            SELECT name, matricule, service
            FROM technicians
            WHERE name LIKE ? OR matricule LIKE ? OR service LIKE ?
            ORDER BY name
            LIMIT 8
            """, (like, like, like))
            rows = cur.fetchall()
            conn.close()

            for name, matricule, service in rows:
                label = f"{name} | {matricule or '-'} | {service or '-'}"

                def choose(n=name):
                    selected_tech["name"] = n
                    tech_entry.delete(0, "end")
                    tech_entry.insert(0, n)
                    clear_suggestions(tech_suggestions)

                ctk.CTkButton(
                    tech_suggestions,
                    text=label,
                    anchor="w",
                    height=32,
                    fg_color="#DDF1FF",
                    text_color="#071B2C",
                    hover_color="#BFE5FF",
                    command=choose
                ).pack(fill="x", padx=5, pady=3)

        def search_tools(event=None):
            clear_suggestions(tool_suggestions)
            selected_tool["name"] = None
            selected_tool["serial"] = None

            keyword = tool_entry.get().strip()
            if not keyword:
                return

            conn = connect_db()
            cur = conn.cursor()
            like = f"%{keyword}%"
            cur.execute("""
            SELECT name, serial, category, status
            FROM tools
            WHERE status='Disponible'
              AND (name LIKE ? OR serial LIKE ? OR category LIKE ?)
            ORDER BY name
            LIMIT 8
            """, (like, like, like))
            rows = cur.fetchall()
            conn.close()

            for name, serial, category, status in rows:
                label = f"{name} | Série: {serial} | {category or '-'} | {status}"

                def choose(n=name, s=serial):
                    selected_tool["name"] = n
                    selected_tool["serial"] = s
                    tool_entry.delete(0, "end")
                    tool_entry.insert(0, f"{n} | {s}")
                    clear_suggestions(tool_suggestions)

                ctk.CTkButton(
                    tool_suggestions,
                    text=label,
                    anchor="w",
                    height=32,
                    fg_color="#DDF1FF",
                    text_color="#071B2C",
                    hover_color="#BFE5FF",
                    command=choose
                ).pack(fill="x", padx=5, pady=3)

        def save_sortie(event=None):
            technician = selected_tech["name"]
            tool_name = selected_tool["name"]
            serial = selected_tool["serial"]

            if not technician:
                messagebox.showwarning("Attention", "Sélectionnez un technicien depuis les résultats.")
                return

            if not tool_name or not serial:
                messagebox.showwarning("Attention", "Sélectionnez un outil disponible depuis les résultats.")
                return

            conn = connect_db()
            cur = conn.cursor()
            cur.execute("UPDATE tools SET status='Sortie' WHERE serial=?", (serial,))
            cur.execute("""
            INSERT INTO movements (technician, tool_name, serial, action, date_time)
            VALUES (?, ?, ?, ?, ?)
            """, (
                technician,
                tool_name,
                serial,
                "Sortie",
                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ))
            conn.commit()
            conn.close()

            messagebox.showinfo("Succès", "Sortie enregistrée avec succès.")
            self.home()

        tech_entry.bind("<KeyRelease>", search_technicians)
        tool_entry.bind("<KeyRelease>", search_tools)
        tech_entry.bind("<Return>", save_sortie)
        tool_entry.bind("<Return>", save_sortie)

        ctk.CTkButton(
            movement_box,
            text="Enregistrer sortie",
            height=40,
            font=("Arial", 15, "bold"),
            fg_color="#2E7D32",
            hover_color="#1B5E20",
            command=save_sortie
        ).grid(row=1, column=2, padx=10, pady=10, sticky="ew")

        active_frame = ctk.CTkScrollableFrame(body, fg_color="white", corner_radius=18)
        active_frame.grid(row=2, column=0, sticky="nsew", pady=10)

        ctk.CTkLabel(
            active_frame,
            text="Mouvements actifs - outils non retournés",
            font=("Arial", 20, "bold"),
            text_color="#071B2C"
        ).pack(anchor="w", padx=15, pady=(15, 8))

        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
        SELECT m.technician, m.tool_name, m.serial, m.date_time
        FROM movements m
        INNER JOIN tools t ON t.serial = m.serial
        WHERE t.status='Sortie'
          AND m.id = (
              SELECT MAX(m2.id)
              FROM movements m2
              WHERE m2.serial = m.serial
          )
          AND m.action='Sortie'
        ORDER BY m.date_time DESC
        """)
        active_rows = cur.fetchall()
        conn.close()

        if not active_rows:
            empty = ctk.CTkFrame(active_frame, fg_color="#F4FAFF", corner_radius=12)
            empty.pack(fill="x", padx=10, pady=6)
            ctk.CTkLabel(
                empty,
                text="Aucun outil sorti actuellement.",
                font=("Arial", 15),
                text_color="#36556B"
            ).pack(anchor="w", padx=15, pady=12)

        for technician, tool_name, serial, date_time in active_rows:
            row = ctk.CTkFrame(active_frame, fg_color="#F4FAFF", corner_radius=12)
            row.pack(fill="x", padx=10, pady=6)

            ctk.CTkLabel(
                row,
                text=f"{date_time} | {technician} | {tool_name} | Série: {serial}",
                font=("Arial", 14),
                text_color="#C62828",
                anchor="w"
            ).pack(side="left", fill="x", expand=True, padx=15, pady=12)

            def return_item(t=technician, n=tool_name, s=serial):
                if messagebox.askyesno("Confirmation", f"Confirmer le retour de l'outil: {n} ?"):
                    conn = connect_db()
                    cur = conn.cursor()
                    cur.execute("UPDATE tools SET status='Disponible' WHERE serial=?", (s,))
                    cur.execute("""
                    INSERT INTO movements (technician, tool_name, serial, action, date_time)
                    VALUES (?, ?, ?, ?, ?)
                    """, (
                        t,
                        n,
                        s,
                        "Retour",
                        datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    ))
                    conn.commit()
                    conn.close()
                    messagebox.showinfo("Succès", "Retour enregistré avec succès.")
                    self.home()

            ctk.CTkButton(
                row,
                text="Retour",
                width=100,
                fg_color="#2E7D32",
                hover_color="#1B5E20",
                command=return_item
            ).pack(side="right", padx=10, pady=8)

    def technicians_page(self):
        self.clear_main()
        self.page_title("Techniciens", "Ajouter, modifier, supprimer et rechercher les techniciens")
        self.form_page("technicians")

    def tools_page(self):
        self.clear_main()
        self.page_title("Outils", "Ajouter, modifier, supprimer et rechercher les outils")
        self.form_page("tools")

    def form_page(self, page_type):
        body = ctk.CTkFrame(self.main, fg_color="#EAF6FF")
        body.grid(row=1, column=0, sticky="nsew", padx=25, pady=10)
        body.grid_columnconfigure(0, weight=1)
        body.grid_rowconfigure(2, weight=1)

        form = self.card(body)
        form.grid(row=0, column=0, sticky="ew", pady=10)
        form.grid_columnconfigure((0, 1, 2), weight=1)

        if page_type == "technicians":
            e1 = ctk.CTkEntry(form, placeholder_text="Nom du technicien", height=40)
            e2 = ctk.CTkEntry(form, placeholder_text="Matricule", height=40)
            e3 = ctk.CTkEntry(form, placeholder_text="Service", height=40)
            search_ph = "Rechercher technicien..."
        else:
            e1 = ctk.CTkEntry(form, placeholder_text="Nom de l'outil", height=40)
            e2 = ctk.CTkEntry(form, placeholder_text="Numéro de série", height=40)
            e3 = ctk.CTkEntry(form, placeholder_text="Catégorie", height=40)
            search_ph = "Rechercher outil..."

        e1.grid(row=0, column=0, padx=10, pady=15, sticky="ew")
        e2.grid(row=0, column=1, padx=10, pady=15, sticky="ew")
        e3.grid(row=0, column=2, padx=10, pady=15, sticky="ew")

        search_entry = ctk.CTkEntry(body, placeholder_text=search_ph, height=42)
        search_entry.grid(row=1, column=0, sticky="ew", pady=8)

        list_frame = ctk.CTkScrollableFrame(body, fg_color="white", corner_radius=18)
        list_frame.grid(row=2, column=0, sticky="nsew", pady=10)

        edit_id = {"value": None}

        def clear_form():
            edit_id["value"] = None
            e1.delete(0, "end")
            e2.delete(0, "end")
            e3.delete(0, "end")
            save_btn.configure(text="Ajouter")

        def refresh():
            for w in list_frame.winfo_children():
                w.destroy()

            keyword = search_entry.get().strip()
            conn = connect_db()
            cur = conn.cursor()

            if page_type == "technicians":
                if keyword:
                    like = f"%{keyword}%"
                    cur.execute("""
                    SELECT id, name, matricule, service
                    FROM technicians
                    WHERE name LIKE ? OR matricule LIKE ? OR service LIKE ?
                    ORDER BY name
                    """, (like, like, like))
                else:
                    cur.execute("SELECT id, name, matricule, service FROM technicians ORDER BY name")
            else:
                if keyword:
                    like = f"%{keyword}%"
                    cur.execute("""
                    SELECT id, name, serial, category, status
                    FROM tools
                    WHERE name LIKE ? OR serial LIKE ? OR category LIKE ? OR status LIKE ?
                    ORDER BY name
                    """, (like, like, like, like))
                else:
                    cur.execute("SELECT id, name, serial, category, status FROM tools ORDER BY name")

            rows = cur.fetchall()
            conn.close()

            for row_data in rows:
                row = ctk.CTkFrame(list_frame, fg_color="#F4FAFF", corner_radius=12)
                row.pack(fill="x", padx=10, pady=6)

                if page_type == "technicians":
                    item_id, a, b, c = row_data
                    text = f"{a} | Matricule: {b or '-'} | Service: {c or '-'}"
                    color = "#071B2C"
                else:
                    item_id, a, b, c, status = row_data
                    text = f"{a} | Série: {b} | Catégorie: {c or '-'} | État: {status}"
                    color = "#2E7D32" if status == "Disponible" else "#C62828"

                ctk.CTkLabel(
                    row,
                    text=text,
                    font=("Arial", 15),
                    text_color=color
                ).pack(side="left", padx=15, pady=12)

                def load_for_edit(data=row_data):
                    edit_id["value"] = data[0]
                    e1.delete(0, "end")
                    e2.delete(0, "end")
                    e3.delete(0, "end")
                    e1.insert(0, data[1] or "")
                    e2.insert(0, data[2] or "")
                    e3.insert(0, data[3] or "")
                    save_btn.configure(text="Modifier")

                def delete_item(id_=item_id):
                    if messagebox.askyesno("Confirmation", "Supprimer cet élément ?"):
                        conn = connect_db()
                        cur = conn.cursor()
                        if page_type == "technicians":
                            cur.execute("DELETE FROM technicians WHERE id=?", (id_,))
                        else:
                            cur.execute("DELETE FROM tools WHERE id=?", (id_,))
                        conn.commit()
                        conn.close()
                        refresh()

                ctk.CTkButton(row, text="Supprimer", width=100, fg_color="#C62828", command=delete_item).pack(side="right", padx=8)
                ctk.CTkButton(row, text="Modifier", width=100, fg_color="#607D8B", command=load_for_edit).pack(side="right", padx=8)

        def save(event=None):
            a = e1.get().strip()
            b = e2.get().strip()
            c = e3.get().strip()

            if page_type == "technicians" and not a:
                messagebox.showwarning("Attention", "Entrez le nom du technicien.")
                return

            if page_type == "tools" and (not a or not b):
                messagebox.showwarning("Attention", "Entrez le nom et le numéro de série.")
                return

            try:
                conn = connect_db()
                cur = conn.cursor()

                if page_type == "technicians":
                    if edit_id["value"]:
                        cur.execute("UPDATE technicians SET name=?, matricule=?, service=? WHERE id=?", (a, b, c, edit_id["value"]))
                    else:
                        cur.execute("INSERT INTO technicians (name, matricule, service) VALUES (?, ?, ?)", (a, b, c))
                else:
                    if edit_id["value"]:
                        cur.execute("UPDATE tools SET name=?, serial=?, category=? WHERE id=?", (a, b, c, edit_id["value"]))
                    else:
                        cur.execute("INSERT INTO tools (name, serial, category) VALUES (?, ?, ?)", (a, b, c))

                conn.commit()
                conn.close()
                clear_form()
                refresh()

            except sqlite3.IntegrityError:
                messagebox.showerror("Erreur", "Donnée déjà existante.")

        save_btn = ctk.CTkButton(form, text="Ajouter", width=130, height=40, command=save)
        save_btn.grid(row=0, column=3, padx=10, pady=15)

        e1.bind("<Return>", save)
        e2.bind("<Return>", save)
        e3.bind("<Return>", save)
        search_entry.bind("<KeyRelease>", lambda e: refresh())

        refresh()

    def journal_page(self):
        self.clear_main()
        self.page_title("Journal", "Historique complet des mouvements")

        body = ctk.CTkFrame(self.main, fg_color="#EAF6FF")
        body.grid(row=1, column=0, sticky="nsew", padx=25, pady=10)
        body.grid_columnconfigure(0, weight=1)
        body.grid_rowconfigure(2, weight=1)

        top_frame = ctk.CTkFrame(body, fg_color="#EAF6FF")
        top_frame.grid(row=0, column=0, sticky="ew", pady=8)
        top_frame.grid_columnconfigure(0, weight=1)

        search_entry = ctk.CTkEntry(top_frame, placeholder_text="Rechercher dans le journal...", height=42)
        search_entry.grid(row=0, column=0, sticky="ew", padx=(0, 8))

        period_box = ctk.CTkComboBox(top_frame, values=["Jour", "Mois", "Année"], width=120, height=42)
        period_box.set("Jour")
        period_box.grid(row=0, column=1, padx=5)

        date_entry = ctk.CTkEntry(top_frame, placeholder_text="2026-05-02 / 2026-05 / 2026", width=220, height=42)
        date_entry.grid(row=0, column=2, padx=5)

        buttons_frame = ctk.CTkFrame(body, fg_color="#EAF6FF")
        buttons_frame.grid(row=1, column=0, sticky="ew", pady=5)

        list_frame = ctk.CTkScrollableFrame(body, fg_color="white", corner_radius=18)
        list_frame.grid(row=2, column=0, sticky="nsew", pady=10)

        def get_rows():
            keyword = search_entry.get().strip()
            conn = connect_db()
            cur = conn.cursor()

            if keyword:
                like = f"%{keyword}%"
                cur.execute("""
                SELECT technician, tool_name, serial, action, date_time
                FROM movements
                WHERE technician LIKE ? OR tool_name LIKE ? OR serial LIKE ? OR action LIKE ? OR date_time LIKE ?
                ORDER BY id DESC
                """, (like, like, like, like, like))
            else:
                cur.execute("""
                SELECT technician, tool_name, serial, action, date_time
                FROM movements
                ORDER BY id DESC
                """)

            rows = cur.fetchall()
            conn.close()
            return rows

        def refresh():
            for w in list_frame.winfo_children():
                w.destroy()

            for technician, tool_name, serial, action, date_time in get_rows():
                row = ctk.CTkFrame(list_frame, fg_color="#F4FAFF", corner_radius=12)
                row.pack(fill="x", padx=10, pady=6)

                color = "#C62828" if action == "Sortie" else "#2E7D32"
                text = f"{date_time} | {action} | {technician} | {tool_name} | Série: {serial}"

                ctk.CTkLabel(
                    row,
                    text=text,
                    font=("Arial", 15),
                    text_color=color,
                    anchor="w"
                ).pack(fill="x", padx=15, pady=12)

        def period_prefix():
            value = date_entry.get().strip()
            if value:
                return value

            now = datetime.now()
            if period_box.get() == "Jour":
                return now.strftime("%Y-%m-%d")
            if period_box.get() == "Mois":
                return now.strftime("%Y-%m")
            return now.strftime("%Y")

        def export_word():
            prefix = period_prefix()
            conn = connect_db()
            cur = conn.cursor()
            cur.execute("""
            SELECT technician, tool_name, serial, action, date_time
            FROM movements
            WHERE date_time LIKE ?
            ORDER BY id DESC
            """, (f"{prefix}%",))
            rows = cur.fetchall()
            conn.close()

            if not rows:
                messagebox.showwarning("Attention", "Aucune donnée pour cette période.")
                return

            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word document", "*.docx")],
                initialfile=f"journal_SMRA_{prefix}.docx"
            )

            if not file_path:
                return

            doc = Document()
            doc.add_heading("Journal Magasin D'outillage SMRA", level=1)
            doc.add_paragraph(f"Période: {period_box.get()} - {prefix}")
            doc.add_paragraph(f"Exporté le: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            headers = table.rows[0].cells
            headers[0].text = "Date"
            headers[1].text = "Action"
            headers[2].text = "Technicien"
            headers[3].text = "Outil"
            headers[4].text = "Série"

            for technician, tool_name, serial, action, date_time in rows:
                cells = table.add_row().cells
                cells[0].text = str(date_time)
                cells[1].text = str(action)
                cells[2].text = str(technician)
                cells[3].text = str(tool_name)
                cells[4].text = str(serial)

            doc.save(file_path)
            messagebox.showinfo("Succès", "Journal exporté en Word avec succès.")

        ctk.CTkButton(
            buttons_frame,
            text="Exporter Word",
            width=170,
            fg_color="#2E7D32",
            hover_color="#1B5E20",
            command=export_word
        ).pack(side="left", padx=5)

        search_entry.bind("<KeyRelease>", lambda e: refresh())
        refresh()

    def settings_page(self):
        self.clear_main()
        self.page_title("Paramètres", "Mot de passe, sauvegarde et restauration")

        body = ctk.CTkFrame(self.main, fg_color="#EAF6FF")
        body.grid(row=1, column=0, sticky="nsew", padx=25, pady=10)

        pass_box = self.card(body)
        pass_box.pack(fill="x", padx=20, pady=15)

        ctk.CTkLabel(
            pass_box,
            text="Changer le mot de passe",
            font=("Arial", 22, "bold"),
            text_color="#071B2C"
        ).pack(anchor="w", padx=25, pady=(25, 10))

        old_pass = ctk.CTkEntry(pass_box, placeholder_text="Ancien mot de passe", show="*", height=42)
        new_pass = ctk.CTkEntry(pass_box, placeholder_text="Nouveau mot de passe", show="*", height=42)
        confirm_pass = ctk.CTkEntry(pass_box, placeholder_text="Confirmer le nouveau mot de passe", show="*", height=42)

        old_pass.pack(fill="x", padx=25, pady=8)
        new_pass.pack(fill="x", padx=25, pady=8)
        confirm_pass.pack(fill="x", padx=25, pady=8)

        def change_password():
            if old_pass.get() != get_password():
                messagebox.showerror("Erreur", "Ancien mot de passe incorrect.")
                return

            if not new_pass.get():
                messagebox.showwarning("Attention", "Entrez un nouveau mot de passe.")
                return

            if new_pass.get() != confirm_pass.get():
                messagebox.showerror("Erreur", "Les mots de passe ne correspondent pas.")
                return

            set_password(new_pass.get())
            old_pass.delete(0, "end")
            new_pass.delete(0, "end")
            confirm_pass.delete(0, "end")
            messagebox.showinfo("Succès", "Mot de passe modifié.")

        ctk.CTkButton(
            pass_box,
            text="Modifier le mot de passe",
            height=42,
            command=change_password
        ).pack(fill="x", padx=25, pady=(10, 25))

        backup_box = self.card(body)
        backup_box.pack(fill="x", padx=20, pady=15)

        ctk.CTkLabel(
            backup_box,
            text="Sauvegarde des données",
            font=("Arial", 22, "bold"),
            text_color="#071B2C"
        ).pack(anchor="w", padx=25, pady=(25, 10))

        def backup_db():
            file_path = filedialog.asksaveasfilename(
                defaultextension=".db",
                filetypes=[("Database", "*.db")],
                initialfile=f"backup_magasin_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"
            )
            if file_path:
                shutil.copy2(DB_NAME, file_path)
                messagebox.showinfo("Succès", "Sauvegarde créée avec succès.")

        def restore_db():
            file_path = filedialog.askopenfilename(filetypes=[("Database", "*.db")])
            if file_path and messagebox.askyesno("Confirmation", "Restaurer cette sauvegarde ?"):
                shutil.copy2(file_path, DB_NAME)
                messagebox.showinfo("Succès", "Sauvegarde restaurée. Redémarrez l'application.")
                self.destroy()

        ctk.CTkButton(
            backup_box,
            text="Créer une sauvegarde",
            height=42,
            fg_color="#2E7D32",
            command=backup_db
        ).pack(fill="x", padx=25, pady=10)

        ctk.CTkButton(
            backup_box,
            text="Restaurer une sauvegarde",
            height=42,
            fg_color="#607D8B",
            command=restore_db
        ).pack(fill="x", padx=25, pady=(5, 25))


if __name__ == "__main__":
    init_db()
    login = Login()
    login.mainloop()
